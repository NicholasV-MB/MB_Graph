import time
from docx import Document
from docx.shared import Inches
from datetime import datetime, timedelta
from MB_Calendar.common_utils import remove_all_files_inside_folder


def write_to_word_doc(year, month, week, start_date, max_days, data):
    """
    Method for write word document with info from @data
    @param: year
    @param: month string with month name
    @param: week number of selected week
    @param: start_date datetime
    @param: max_days
    @param: data list of dict with info to write

    @return: file_name: string with path inside static folder
    """
    current_dir = "MB_Calendar/static/"
    # first of all: remove all files from temp_files folder
    mypath = current_dir+"temp_files/"
    remove_all_files_inside_folder(mypath)

    file_name = "temp_files/PLAN_WEEK"+str(week)+'_OF_'+month+'_'+ year+"_MAX"+max_days+"DAYS.docx"
    document = Document()

    document.add_heading('Plan for Week '+str(week)+' of '+month+' '+ year, 0)

    p = document.add_paragraph('First Day : ')
    p.add_run(start_date.strftime("%d.%m.%y")).bold = True
    p = document.add_paragraph('Last Day : ')
    last_day = start_date + timedelta(4)
    p.add_run(last_day.strftime("%d.%m.%y")).bold = True

    for activity in data:
        time_delta = int(activity["day"])-1
        new_date = start_date + timedelta(time_delta)
        document.add_heading(new_date.strftime("%d.%m.%y"), level=1)
        document.add_paragraph(activity["description"])
        document.add_paragraph(activity["location"])
        document.add_paragraph("Duration: "+str(activity["duration"])+" Hours")
        document.add_paragraph(activity["info"])

    document.save(current_dir+file_name)

    return file_name
